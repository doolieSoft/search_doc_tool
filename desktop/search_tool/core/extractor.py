import re

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import fitz
    HAS_PDF = True
except ImportError:
    HAS_PDF = False


def extract_text_docx(path: str) -> str:
    doc = docx.Document(path)
    texts = []

    def extract_para_text(para) -> str:
        # Concaténer les runs SANS espace — Word fragmente parfois les mots
        raw = "".join(run.text for run in para.runs)
        return re.sub(r'\s+', ' ', raw).strip()

    def extract_element(elem):
        for para in elem.paragraphs:
            t = extract_para_text(para)
            if t:
                texts.append(t)
        for table in elem.tables:
            for row in table.rows:
                for cell in row.cells:
                    extract_element(cell)

    extract_element(doc)
    return " ".join(texts)


def extract_text_pdf(path: str) -> list[tuple[int, str]]:
    """Retourne liste de (page_num, texte) pour chaque page."""
    doc = fitz.open(path)
    pages = [(i + 1, re.sub(r'\s+', ' ', page.get_text())) for i, page in enumerate(doc)]
    doc.close()
    return pages

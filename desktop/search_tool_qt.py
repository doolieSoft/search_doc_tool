"""
Outil de recherche de termes dans fichiers Word (.docx) et PDF
Interface PyQt6 — dark mode
Dépendances : pip install PyQt6 python-docx pymupdf
"""

import os
import re
import json
import time
import unicodedata
import threading

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QCheckBox, QFileDialog,
    QTableView, QHeaderView, QAbstractItemView, QStatusBar,
    QStyledItemDelegate, QStyleOptionViewItem, QSizePolicy,
    QFrame, QProgressBar,
)
from PyQt6.QtCore import (
    Qt, QThread, pyqtSignal, QAbstractTableModel, QModelIndex,
    QVariant, QRect, QRectF, QSize,
)
from PyQt6.QtWidgets import QStyle
from PyQt6.QtGui import (
    QFont, QColor, QPainter, QTextDocument, QTextOption,
    QPalette, QTextCharFormat, QTextCursor, QAbstractTextDocumentLayout,
)

# ── Imports optionnels ──────────────────────────────────────────────────────
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import fitz
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

ICON_PNG_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAQAAAAEACAYAAABccqhmAAAIVElEQVR4nO3dyXUbRxRA0aaPA1EWXjgDh6EEFIsTcBjKQAtnoUzshcQjECRIDDX84d4AJKKr/utqkASPAwAAAAAAAAAAAAAAAAAAAAAAAADY4Wn3F8APf/3z33+j/q2vn5+sK1exURYaOeT3EgdO2QwTRRj4jwhCbxZ/oAwD/xFB6MViP6jC0F8iBvVZ4DtUHvpLxKAmi3qDjoN/TghqsZgfMPSXiUF+FvACg389IcjLwp0x+PcTgnws2E8GfxwhyKP9Qhn8eYQgvt92fwE7Gf65XN/4WhbaxlzPaSCmVoti8PcTgljaPAIY/hisQywtAmDTxWI94ih9HLPR4vNIsFfZE4Dhz8E67VUyADZVLtZrn1LHr+wb6fvffz78b3z68m3AV7KPR4K1ylzsTMM/YtBvlSkMIrBOiQsdffh3DPxHogdBBNZIf5GjDn/Eob8kagxEYL7UFzja8Gca+kuixUAE5kp7cSMNf4XBPxcpBCIwT8oLG2H4Kw79JRFiIAJzpLuou4e/0+Cf2x0CERgv1Q8CGf69dr/+3etfUZqi7lz83Rs/op2nASeBcVJcyF3Db/A/tisEIjBG+EcAwx/bruvkcWCM33d/AdEY/Ns9X7PdbxJyu9AngNWVN/yPWX39nAIeFzYAhj8nEcgl5BspKxfV4M+z8pHAm4L3CXsCWMHwz+X6xhcuAKvu/jbnGquus0eB+4QKgOGvSQTiChMAw1+bCMQUJgArGP69XP94QgRgRbVtvhhWrINTwPVCBGA2wx+L9YhjewBm19pmi2n2ujgFXGdrAAx/byKw3/YTALDPtgC4+3McTgG7lTwBGP5crNc+WwIws8o2U04z180p4LKSJwDgOssD4O7PJU4B65U5ARj+GqzjWksDoMLsZP+9VuIE4K5Ri/Vcp0QAgPssC8Cs45e7RU2z1tVjwEupTwCGvzbrO9+SAKgukdiPv6Q9Abg79GCd50obAOBx0wPguEVE9uUPKU8AjoW9WO95UgYAGGPq31ObcczadTf49OXbv1v+46C+//3nH6v/zxl/a7D73xR0AoDGBAAaSxUAbwb1Zv3HmxYA32Yhg+77NNUJABgrTQAc/zgO+2C0NAEAxpsSgO7PVeTSeb86AUBjKQLguY9T9sM4KQIAzCEA0Njvu7+ALHb88gvM5gQAjQ0PQOdvqZBX130b/gTgHV/eYl+MET4AwDwCAI0JADQmANCYnwO4kg8FfcnPRdTgBACNCQA0JgDQmABAYwIAjQkANCYA0JgAQGMCAI0JADQmANBY+AB8+vJt95dAQPbFGE8z/tHRH6/k0184NzoAXz8/TZmF6MKfAIB5BAAaEwBoTACgsRQB8I4vp+yHcVIEAJhjSgC6fkuFnDrvVx8KeqUMHwrqgzq5VZpHAM99HId9MFqaAADjTQtA5+cq8ui+T1OdABz/erP+46UKADCWAEBjUwMw4/nKMbCnGeve/fn/OJwAoLWUAXAK6MV6zzM9AI5ZRGRf/pDyBACMkTYAjoU9WOe5lh2DRn9Q6HH4sNAOvPs/V9oTwHG4O1RnfedbFoBZ1bVJapq1ru7+L6U+AQCPKREAp4BarOc6SwPg+MVO9t9rJU4Ax+GuUYV1XGt5AGZW2ObJbeb6ufu/rcwJALjdlgA4BXDO3X+PkicAEcjFeu2zLQCzq2xT5TB7ndz931fyBABcZ2sAnAJ6c/ffb/sJQAR6MvwxbA/ACiIQi/WII0QAVtTapothxTq4+18vRABWEYG9XP94wgRgVbVtwj1WXXd3/9uECcBxiEBVhj+uUAE4DhGoxvDHFi4AK4nAXK5vfGGrOeNThN/jE4bHWT347v73C3sCWL2o7lZjGP5cwgbgOEQgG9cvn/D1XP0o8MwjwfV2D75TwP1CnwCOY9/i7t7UWUS4TrtuEhWkKefORXYaeC3C4J9zErhdqgu2u/RCEHPwT4nAbcI/ApzavbjRN/9sGV7/7ptENilrGWGRO50GMgz+ud03iyzSXqQIEXhWMQYZh/6cCHws9QWKFIHjqBGCCoN/SgTel/7iRIvAs0wxqDb050TgshIXJmoEnkWMQfWhPycCbytzUaJH4NSOIHQb+LeIwGulLkimCLxlRBgM+vtE4KWSFyN7CJhLBH5J9YNA17LAuXz9/PS0cs3cIH4pGYDjEIEsTtdJBNZrMSQWO573hn3lenW/UZQ9AZzqvsjRfLQeTgLrtAjAcYhAFNeugwis0XIoOi/4LvcOtMeBudq94FNCMN+IoRKBedo8Aryl22KvNur6ehyYxwD81G3hZ5o1sE4C47V4kbcQgvutGBoRGKv8C7yXEFxv9aCIwDilX9wIQnDZzuEQgTHKvrAZxCDWMIjA40q+qNk6hiDqAIjAY8q9oNUqxyDLhheB+5V6MbtViEHWDS4C9ynzQiLKEIRKm1kEblfiRWQRIQhVNu4lInCb9C+gipEbt8LGfIQIXC/1Fw+XiMB1Wv8yEHX5BaLrCABlicDHBIDSROB9AkB5InCZANCCCLxNAGhDBF4TAFoRgZcEgHZE4BcBoCUR+EEAaEsEBIDmukdAAGivcwQEAI6+ERAA+KljBAQATnSLgADAmU4REAB4Q5cICABc0CECAgDvqB6BtJ9lBiutHs5V4XECgCtU/QvIAgBXqhgBAYAbVIuAAMCNKkVAAOAOmf8YyCkBgDtV+BahAMADsp8EBAAelDkCAgADZI2AAMAgGSMgADBQtggIAAyWKQICABOMjsCsqAgANCYAMMmou/bMRwoBgIkeHd7Z7ycIAEx27xCveDNRAGCBW4d51XcS0ny7Aqp47xd7Mn0LEQAAAAAAAAAAAAAAAAAAAAAAAAAAuMP/9Qdy5QvdFDgAAAAASUVORK5CYII="
)

CONFIG_FILE = os.path.join(os.path.expanduser("~"), ".search_tool_config.json")

# ── Palette dark ────────────────────────────────────────────────────────────
BG_DARK      = "#1A1D23"
BG_PANEL     = "#22262F"
BG_INPUT     = "#2A2F3A"
BG_ROW_ODD   = "#22262F"
BG_ROW_EVEN  = "#1E2229"
BG_ROW_SEL   = "#1E3A5F"
BORDER       = "#333844"
TEXT_PRIMARY = "#E8EAF0"
TEXT_MUTED   = "#7B8196"
ACCENT       = "#4A9EFF"
ACCENT_DARK  = "#1A5FA8"
HIGHLIGHT    = "#FFB347"   # couleur des termes trouvés dans le contexte
GREEN        = "#3DAA6E"
RED          = "#E05252"

STYLESHEET = f"""
QMainWindow, QWidget {{
    background-color: {BG_DARK};
    color: {TEXT_PRIMARY};
    font-family: 'Segoe UI', 'Ubuntu', sans-serif;
    font-size: 13px;
}}
QLabel {{
    color: {TEXT_PRIMARY};
}}
QLineEdit {{
    background-color: {BG_INPUT};
    border: 1px solid {BORDER};
    border-radius: 6px;
    padding: 6px 10px;
    color: {TEXT_PRIMARY};
    selection-background-color: {ACCENT_DARK};
}}
QLineEdit:focus {{
    border: 1px solid {ACCENT};
}}
QPushButton {{
    background-color: {BG_INPUT};
    border: 1px solid {BORDER};
    border-radius: 6px;
    padding: 6px 14px;
    color: {TEXT_PRIMARY};
}}
QPushButton:hover {{
    background-color: {BORDER};
}}
QPushButton#btn_search {{
    background-color: {GREEN};
    border: none;
    color: #fff;
    font-weight: bold;
    padding: 8px 20px;
}}
QPushButton#btn_search:hover {{
    background-color: #35956A;
}}
QPushButton#btn_search:disabled {{
    background-color: #2A4A3A;
    color: #668866;
}}
QPushButton#btn_stop {{
    background-color: {RED};
    border: none;
    color: #fff;
    font-weight: bold;
    padding: 8px 20px;
}}
QPushButton#btn_stop:hover {{
    background-color: #C04040;
}}
QPushButton#btn_stop:disabled {{
    background-color: #3A2222;
    color: #885555;
}}
QPushButton#btn_browse {{
    background-color: {ACCENT_DARK};
    border: none;
    color: #fff;
    padding: 6px 14px;
}}
QPushButton#btn_browse:hover {{
    background-color: {ACCENT};
}}
QCheckBox {{
    color: {TEXT_MUTED};
    spacing: 6px;
}}
QCheckBox::indicator {{
    width: 16px; height: 16px;
    border: 1px solid {BORDER};
    border-radius: 4px;
    background: {BG_INPUT};
}}
QCheckBox::indicator:checked {{
    background: {ACCENT};
    border-color: {ACCENT};
}}
QTableView {{
    background-color: {BG_PANEL};
    border: 1px solid {BORDER};
    border-radius: 8px;
    gridline-color: {BORDER};
    selection-background-color: {BG_ROW_SEL};
    outline: none;
}}
QTableView::item {{
    padding: 4px 8px;
    border: none;
}}
QTableView::item:selected {{
    background-color: {BG_ROW_SEL};
    color: {TEXT_PRIMARY};
}}
QHeaderView::section {{
    background-color: {BG_INPUT};
    color: {TEXT_MUTED};
    border: none;
    border-right: 1px solid {BORDER};
    border-bottom: 1px solid {BORDER};
    padding: 6px 10px;
    font-weight: bold;
    font-size: 12px;
    text-transform: uppercase;
    letter-spacing: 1px;
}}
QScrollBar:vertical {{
    background: {BG_DARK};
    width: 8px;
    border-radius: 4px;
}}
QScrollBar::handle:vertical {{
    background: {BORDER};
    border-radius: 4px;
    min-height: 30px;
}}
QScrollBar::handle:vertical:hover {{
    background: {TEXT_MUTED};
}}
QScrollBar:horizontal {{
    background: {BG_DARK};
    height: 8px;
    border-radius: 4px;
}}
QScrollBar::handle:horizontal {{
    background: {BORDER};
    border-radius: 4px;
}}
QStatusBar {{
    color: {TEXT_MUTED};
    background: {BG_DARK};
    font-size: 12px;
}}
QProgressBar {{
    background-color: {BG_INPUT};
    border: none;
    border-radius: 3px;
    height: 4px;
    text-align: center;
}}
QProgressBar::chunk {{
    background-color: {ACCENT};
    border-radius: 3px;
}}
QFrame#separator {{
    background-color: {BORDER};
    max-height: 1px;
}}
"""

# ── Config ───────────────────────────────────────────────────────────────────
def load_config() -> dict:
    try:
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

def save_config(data: dict):
    try:
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def load_favorites() -> list:
    """Retourne liste de dicts {path, name}."""
    favs = load_config().get("favorites", [])
    # Migration : anciens favoris en strings → dicts
    return [f if isinstance(f, dict) else {"path": f, "name": os.path.basename(f)}
            for f in favs]

def save_favorites(favs: list):
    cfg = load_config()
    cfg["favorites"] = favs
    save_config(cfg)

# ── Normalisation ────────────────────────────────────────────────────────────
def remove_accents(s: str) -> str:
    # Normaliser les apostrophes typographiques → apostrophe droite
    s = s.replace(‘’’, "’").replace(‘’’, "’").replace(‘ʼ’, "’")
    # Espaces insécables → espace normale
    s = s.replace(' ', ' ').replace(' ', ' ')
    # Supprimer les diacritiques
    return ''.join(c for c in unicodedata.normalize('NFD', s)
                   if unicodedata.category(c) != 'Mn')

# ── Parser de requête ────────────────────────────────────────────────────────
def parse_query(raw: str):
    raw_no_quotes = re.sub(r'"[^"]+"', '', raw)
    mode = "AND" if '+' in raw_no_quotes else "OR"
    term_list = []
    i = 0
    while i < len(raw):
        if raw[i] == '"':
            end = raw.find('"', i + 1)
            if end != -1:
                phrase = raw[i + 1:end].strip()
                if phrase:
                    term_list.append(phrase)
                i = end + 1
            else:
                i += 1
        elif raw[i] in ('+', ',', ' '):
            i += 1
        else:
            end = i
            while end < len(raw) and raw[end] not in ('"', '+', ',', ' '):
                end += 1
            word = raw[i:end].strip()
            if word:
                term_list.append(word)
            i = end
    return term_list, mode

# ── Extraction de texte ──────────────────────────────────────────────────────
def extract_text_docx(path: str) -> str:
    doc = docx.Document(path)
    texts = []

    def extract_para_text(para) -> str:
        # Concaténer les runs SANS espace — Word fragmente parfois les mots
        # ex: "P" + "XPLO" → "PXPLO" et non "P XPLO"
        raw = "".join(run.text for run in para.runs)
        return re.sub(r'\s+', ' ', raw).strip()

    def extract_element(elem):
        for para in elem.paragraphs:
            t = extract_para_text(para)
            if t:
                texts.append(t)
        for table in elem.tables:
            for row in table.rows:
                for cell in row.cells:
                    extract_element(cell)

    extract_element(doc)
    return " ".join(texts)

def extract_text_pdf(path: str) -> list[tuple[int, str]]:
    """Retourne liste de (page_num, texte) pour chaque page."""
    doc = fitz.open(path)
    pages = [(i + 1, re.sub(r'\s+', ' ', page.get_text())) for i, page in enumerate(doc)]
    doc.close()
    return pages

# ── Contexte ─────────────────────────────────────────────────────────────────
def _word_span(text: str, ts: int, te: int) -> tuple[int, int]:
    """Étend te jusqu'à la fin du mot (lettres, chiffres, _ et -)."""
    while te < len(text) and (text[te].isalnum() or text[te] in '_-'):
        te += 1
    return ts, te

def get_context(text: str, match, window: int = 100) -> str:
    start = max(0, match.start() - window)
    end = min(len(text), match.end() + window)
    snippet = text[start:end].replace("\n", " ").strip()
    ts = match.start() - start
    te = match.end() - start
    ts, te = _word_span(snippet, ts, te)
    snippet = snippet[:ts] + "[" + snippet[ts:te] + "]" + snippet[te:]
    return ("…" if start > 0 else "") + snippet + ("…" if end < len(text) else "")

def get_combined_context(text: str, matches: list, window: int = 100,
                          proximity_words: int = 30) -> str:
    if not matches:
        return ""
    if len(matches) == 1:
        return get_context(text, matches[0], window)
    matches = sorted(matches, key=lambda m: m.start())
    between = text[matches[0].end():matches[-1].start()]
    if len(between.split()) <= proximity_words:
        start = max(0, matches[0].start() - window)
        end = min(len(text), matches[-1].end() + window)
        snippet = text[start:end].replace("\n", " ").strip()
        offset = start
        for m in reversed(matches):
            ts = m.start() - offset
            te = m.end() - offset
            ts, te = _word_span(snippet, ts, te)
            snippet = snippet[:ts] + "[" + snippet[ts:te] + "]" + snippet[te:]
        return ("…" if start > 0 else "") + snippet + ("…" if end < len(text) else "")
    else:
        return " | ".join(get_context(text, m, window) for m in matches)

# ── Recherche ────────────────────────────────────────────────────────────────
def build_pattern(term: str, case_sensitive: bool, whole_word: bool):
    norm = remove_accents(term) if not case_sensitive else term
    words = re.split(r'\s+', norm.strip())
    escaped = [re.escape(w) for w in words]
    pattern_str = r'[\s\W]*'.join(escaped)
    if whole_word:
        pattern_str = r'\b' + pattern_str + r'\b'
    flags = 0 if case_sensitive else re.IGNORECASE
    return re.compile(pattern_str, flags)

def search_file(path: str, terms: list, case_sensitive: bool,
                whole_word: bool, mode: str) -> list[dict]:
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx":
            if not HAS_DOCX:
                return []
            raw_text = extract_text_docx(path)
            # DOCX : un seul bloc, pas de page
            pages = [(None, raw_text)]
        elif ext == ".pdf":
            if not HAS_PDF:
                return []
            pages = extract_text_pdf(path)
        else:
            return []
    except Exception as e:
        return [{"file": path, "term": "ERREUR", "context": str(e), "page": None}]

    if mode == "AND":
        # Chercher sur le texte complet (concaténé pour docx, par page pour pdf)
        results = []
        for page_num, raw_text in pages:
            search_text = remove_accents(raw_text) if not case_sensitive else raw_text
            patterns = {}
            for term in terms:
                try:
                    patterns[term] = build_pattern(term, case_sensitive, whole_word)
                except re.error:
                    return []
            page_matches = {}
            all_found = True
            for term, pat in patterns.items():
                ms = list(pat.finditer(search_text))
                if not ms:
                    all_found = False
                    break
                page_matches[term] = ms
            if all_found:
                best = [ms[0] for ms in page_matches.values()]
                # Extrait brut autour du premier match pour Ctrl+F
                m0 = best[0]
                s = max(0, m0.start() - 20)
                e = min(len(raw_text), m0.end() + 20)
                ctrlf = raw_text[s:e].replace("\n", " ").strip()
                results.append({"file": path, "term": " + ".join(terms),
                                 "context": get_combined_context(raw_text, best),
                                 "ctrlf": ctrlf,
                                 "page": page_num})
        return results
    else:
        results = []
        for page_num, raw_text in pages:
            search_text = remove_accents(raw_text) if not case_sensitive else raw_text
            for term in terms:
                try:
                    pat = build_pattern(term, case_sensitive, whole_word)
                except re.error:
                    continue
                for match in pat.finditer(search_text):
                    # Extrait brut ~20 chars autour du match pour Ctrl+F Word
                    s = max(0, match.start() - 20)
                    e = min(len(raw_text), match.end() + 20)
                    ctrlf = raw_text[s:e].replace("\n", " ").strip()
                    results.append({"file": path, "term": term,
                                     "context": get_context(raw_text, match),
                                     "ctrlf": ctrlf,
                                     "page": page_num})
        return results

def collect_files(folder: str, recurse: bool) -> list:
    exts = {".docx", ".pdf"}
    files = []
    if recurse:
        for root, _, filenames in os.walk(folder):
            for f in filenames:
                if os.path.splitext(f)[1].lower() in exts and not f.startswith(("~$", "._", ".~")):
                    files.append(os.path.join(root, f))
    else:
        for f in os.listdir(folder):
            if os.path.splitext(f)[1].lower() in exts and not f.startswith(("~$", "._", ".~")):
                files.append(os.path.join(folder, f))
    return files

# ── Thread de recherche ───────────────────────────────────────────────────────
class SearchWorker(QThread):
    result_found  = pyqtSignal(dict)
    progress      = pyqtSignal(int, int, str)
    finished      = pyqtSignal(int, int)

    def __init__(self, folder, terms, case_sensitive, whole_word, recurse, mode):
        super().__init__()
        self.folder = folder
        self.terms = terms
        self.case_sensitive = case_sensitive
        self.whole_word = whole_word
        self.recurse = recurse
        self.mode = mode
        self._stop = False

    def stop(self):
        self._stop = True

    def run(self):
        from concurrent.futures import ThreadPoolExecutor, as_completed
        files = collect_files(self.folder, self.recurse)
        total = len(files)
        count = 0
        # Nombre de workers : min(8, nb_cœurs, nb_fichiers)
        workers = min(8, os.cpu_count() or 4, max(1, total))

        with ThreadPoolExecutor(max_workers=workers) as executor:
            futures = {
                executor.submit(
                    search_file, path, self.terms,
                    self.case_sensitive, self.whole_word, self.mode
                ): path
                for path in files
            }
            done = 0
            for future in as_completed(futures):
                if self._stop:
                    executor.shutdown(wait=False, cancel_futures=True)
                    break
                done += 1
                path = futures[future]
                self.progress.emit(done, total, os.path.basename(path))
                try:
                    results = future.result()
                except Exception as e:
                    results = [{"file": path, "term": "ERREUR",
                                "context": str(e), "page": None, "ctrlf": ""}]
                for r in results:
                    count += 1
                    self.result_found.emit(r)

        self.finished.emit(count, total)

# ── Modèle de données ─────────────────────────────────────────────────────────
class ResultsModel(QAbstractTableModel):
    HEADERS = ["Fichier", "Terme", "Contexte"]

    def __init__(self):
        super().__init__()
        self._data: list[dict] = []

    def rowCount(self, parent=QModelIndex()):
        return len(self._data)

    def columnCount(self, parent=QModelIndex()):
        return 3

    def data(self, index, role=Qt.ItemDataRole.DisplayRole):
        if not index.isValid() or index.row() >= len(self._data):
            return QVariant()
        row = self._data[index.row()]
        col = index.column()
        if role == Qt.ItemDataRole.DisplayRole:
            return [os.path.basename(row["file"]), row["term"], row["context"]][col]
        if role == Qt.ItemDataRole.ToolTipRole:
            if col == 0:
                return row["file"]
            if col == 2:
                return row["context"]
        if role == Qt.ItemDataRole.UserRole:
            return row
        if role == Qt.ItemDataRole.BackgroundRole:
            if row.get("term") == "ERREUR":
                return QColor("#3A1A1A")
            return QColor(BG_ROW_ODD if index.row() % 2 == 0 else BG_ROW_EVEN)
        if role == Qt.ItemDataRole.ForegroundRole:
            if row.get("term") == "ERREUR":
                return QColor(RED)
            return QColor(TEXT_PRIMARY)
        return QVariant()

    def headerData(self, section, orientation, role=Qt.ItemDataRole.DisplayRole):
        if orientation == Qt.Orientation.Horizontal and role == Qt.ItemDataRole.DisplayRole:
            return self.HEADERS[section]
        return QVariant()

    def add_result(self, r: dict):
        row = len(self._data)
        self.beginInsertRows(QModelIndex(), row, row)
        self._data.append(r)
        self.endInsertRows()

    def clear(self):
        self.beginResetModel()
        self._data = []
        self.endResetModel()

    def get_row(self, idx: int) -> dict:
        return self._data[idx] if idx < len(self._data) else {}

    def all_results(self):
        return list(self._data)

# ── Delegate pour rendu HTML du contexte ─────────────────────────────────────
class ContextDelegate(QStyledItemDelegate):
    """Colonne Contexte : affiche les [termes] en gras et en couleur."""

    def _make_doc(self, text: str, font: QFont, width: int) -> QTextDocument:
        doc = QTextDocument()
        doc.setDefaultFont(font)
        doc.setTextWidth(width)
        # Convertir [terme] en HTML gras coloré
        html = re.sub(
            r'\[([^\]]+)\]',
            rf'<b><span style="color:{HIGHLIGHT};">\1</span></b>',
            re.sub(r'&', '&amp;', re.sub(r'<', '&lt;', text))
        )
        doc.setHtml(f'<span style="color:{TEXT_PRIMARY};">{html}</span>')
        return doc

    def paint(self, painter: QPainter, option: QStyleOptionViewItem, index: QModelIndex):
        if index.column() != 2:
            super().paint(painter, option, index)
            return

        painter.save()

        # Fond
        bg = index.data(Qt.ItemDataRole.BackgroundRole)
        is_selected = bool(option.state & QStyle.StateFlag.State_Selected)
        if is_selected:
            painter.fillRect(option.rect, QColor(BG_ROW_SEL))
        else:
            painter.fillRect(option.rect, bg or QColor(BG_ROW_ODD))

        text = index.data(Qt.ItemDataRole.DisplayRole) or ""
        doc = self._make_doc(text, option.font, option.rect.width() - 16)

        painter.translate(option.rect.left() + 8, option.rect.top() + 4)
        clip = QRectF(0, 0, option.rect.width() - 16, option.rect.height() - 8)
        doc.drawContents(painter, clip)
        painter.restore()

    def sizeHint(self, option: QStyleOptionViewItem, index: QModelIndex) -> QSize:
        if index.column() != 2:
            return super().sizeHint(option, index)
        text = index.data(Qt.ItemDataRole.DisplayRole) or ""
        doc = self._make_doc(text, option.font, option.rect.width() - 16 or 500)
        return QSize(int(doc.idealWidth()) + 16, min(int(doc.size().height()) + 8, 80))


# ── Fenêtre principale ────────────────────────────────────────────────────────
class SearchApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Recherche dans Word / PDF")
        self.resize(1200, 750)
        self._worker: SearchWorker | None = None
        self._model = ResultsModel()
        self._build_ui()
        self._load_config()
        self._set_icon()

    def _set_icon(self):
        import base64
        from PyQt6.QtGui import QPixmap, QIcon
        from PyQt6.QtCore import QByteArray
        data = QByteArray(base64.b64decode(ICON_PNG_B64))
        pix = QPixmap()
        pix.loadFromData(data, "PNG")
        self.setWindowIcon(QIcon(pix))
        QApplication.instance().setWindowIcon(QIcon(pix))

    def _build_ui(self):
        central = QWidget()
        self.setCentralWidget(central)
        root = QVBoxLayout(central)
        root.setContentsMargins(16, 16, 16, 12)
        root.setSpacing(10)

        # ── Dossier ──────────────────────────────────────────────────────────
        row_folder = QHBoxLayout()
        row_folder.setSpacing(8)
        lbl_folder = QLabel("Dossier")
        lbl_folder.setFixedWidth(56)
        lbl_folder.setStyleSheet(f"color:{TEXT_MUTED}; font-size:12px;")
        self.inp_folder = QLineEdit()
        self.inp_folder.setPlaceholderText("Chemin du dossier à analyser…")
        btn_browse = QPushButton("Parcourir")
        btn_browse.setObjectName("btn_browse")
        btn_browse.setFixedWidth(100)
        btn_browse.clicked.connect(self._browse)
        row_folder.addWidget(lbl_folder)
        row_folder.addWidget(self.inp_folder)
        row_folder.addWidget(btn_browse)
        btn_add_fav = QPushButton("★")
        btn_add_fav.setToolTip("Ajouter ce dossier aux favoris")
        btn_add_fav.setFixedWidth(32)
        btn_add_fav.setStyleSheet(f"color: {HIGHLIGHT}; font-size:16px; border:1px solid {BORDER}; border-radius:6px; background:{BG_INPUT};")
        btn_add_fav.clicked.connect(self._add_favorite)
        row_folder.addWidget(btn_add_fav)
        root.addLayout(row_folder)

        # ── Favoris ───────────────────────────────────────────────────────────
        self.fav_frame = QWidget()
        self.fav_frame.setStyleSheet(f"background: transparent;")
        self.fav_layout = QHBoxLayout(self.fav_frame)
        self.fav_layout.setContentsMargins(0, 0, 0, 0)
        self.fav_layout.setSpacing(6)
        fav_label = QLabel("Favoris :")
        fav_label.setStyleSheet(f"color:{TEXT_MUTED}; font-size:12px;")
        fav_label.setFixedWidth(56)
        self.fav_layout.addWidget(fav_label)
        self.fav_layout.addStretch()
        root.addWidget(self.fav_frame)
        self._refresh_favorites()

        # ── Termes ───────────────────────────────────────────────────────────
        row_terms = QHBoxLayout()
        row_terms.setSpacing(8)
        lbl_terms = QLabel("Termes")
        lbl_terms.setFixedWidth(56)
        lbl_terms.setStyleSheet(f"color:{TEXT_MUTED}; font-size:12px;")
        self.inp_terms = QLineEdit()
        self.inp_terms.setPlaceholderText('espace = OR   +  = AND   "..." = expression exacte')
        self.inp_terms.returnPressed.connect(self._start_search)
        hint = QLabel('espace=OR  +  =AND  "…"=exact')
        hint.setStyleSheet(f"color:{TEXT_MUTED}; font-size:11px;")
        hint.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)
        row_terms.addWidget(lbl_terms)
        row_terms.addWidget(self.inp_terms)
        row_terms.addWidget(hint)
        root.addLayout(row_terms)

        # ── Options ───────────────────────────────────────────────────────────
        row_opts = QHBoxLayout()
        row_opts.setSpacing(20)
        self.chk_case      = QCheckBox("Respecter la casse")
        self.chk_word      = QCheckBox("Mot entier")
        self.chk_recurse   = QCheckBox("Sous-dossiers")
        self.chk_recurse.setChecked(True)
        for chk in (self.chk_case, self.chk_word, self.chk_recurse):
            row_opts.addWidget(chk)
        row_opts.addStretch()
        root.addLayout(row_opts)

        # ── Boutons ───────────────────────────────────────────────────────────
        row_btns = QHBoxLayout()
        row_btns.setSpacing(10)
        self.btn_search = QPushButton("🔍  Lancer la recherche")
        self.btn_search.setObjectName("btn_search")
        self.btn_search.clicked.connect(self._start_search)
        self.btn_stop = QPushButton("⏹  Stopper")
        self.btn_stop.setObjectName("btn_stop")
        self.btn_stop.setEnabled(False)
        self.btn_stop.clicked.connect(self._stop_search)
        btn_clear = QPushButton("Effacer")
        btn_clear.clicked.connect(self._clear)
        btn_export = QPushButton("Exporter CSV")
        btn_export.clicked.connect(self._export_csv)
        for b in (self.btn_search, self.btn_stop, btn_clear, btn_export):
            row_btns.addWidget(b)
        row_btns.addStretch()
        root.addLayout(row_btns)

        # ── Barre de progression ──────────────────────────────────────────────
        self.progress = QProgressBar()
        self.progress.setFixedHeight(4)
        self.progress.setTextVisible(False)
        self.progress.hide()
        root.addWidget(self.progress)

        # ── Séparateur ────────────────────────────────────────────────────────
        sep = QFrame()
        sep.setObjectName("separator")
        sep.setFrameShape(QFrame.Shape.HLine)
        root.addWidget(sep)

        # ── Tableau ───────────────────────────────────────────────────────────
        self.table = QTableView()
        self.table.setModel(self._model)
        self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.SelectionMode.SingleSelection)
        self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        self.table.setShowGrid(False)
        self.table.setAlternatingRowColors(False)
        self.table.verticalHeader().hide()
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Interactive)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Interactive)
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        self.table.setColumnWidth(0, 300)
        self.table.setColumnWidth(1, 130)
        self.table.setItemDelegateForColumn(2, ContextDelegate())
        self.table.doubleClicked.connect(self._open_file)
        self.table.setWordWrap(True)
        self.table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        root.addWidget(self.table)

        # ── Status bar ────────────────────────────────────────────────────────
        self.status = QStatusBar()
        self.setStatusBar(self.status)
        self.status.showMessage("Prêt.")

    # ── Config ───────────────────────────────────────────────────────────────
    def _load_config(self):
        cfg = load_config()
        if cfg.get("folder"):
            self.inp_folder.setText(cfg["folder"])

    # ── Favoris ──────────────────────────────────────────────────────────
    def _add_favorite(self):
        folder = self.inp_folder.text().strip()
        if not folder or not os.path.isdir(folder):
            self.status.showMessage("⚠  Dossier invalide.")
            return
        favs = load_favorites()
        if not any(f["path"] == folder for f in favs):
            name = os.path.basename(folder) or folder
            favs.append({"path": folder, "name": name})
            save_favorites(favs)
            self._refresh_favorites()
            self.status.showMessage(f"Favori ajouté : {name}")
        else:
            self.status.showMessage("Ce dossier est déjà dans les favoris.")

    def _refresh_favorites(self):
        while self.fav_layout.count() > 2:
            item = self.fav_layout.takeAt(1)
            if item.widget():
                item.widget().deleteLater()
        favs = load_favorites()
        for fav in favs:
            name = fav["name"]
            path = fav["path"]
            btn = QPushButton(f"📁 {name}")
            btn.setToolTip(path)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {BG_INPUT}; border: 1px solid {BORDER};
                    border-radius: 5px; padding: 3px 10px;
                    color: {TEXT_PRIMARY}; font-size: 12px;
                }}
                QPushButton:hover {{ background: {BORDER}; }}
            """)
            btn.clicked.connect(lambda checked, p=path: self._use_favorite(p))
            btn.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
            btn.customContextMenuRequested.connect(
                lambda pos, f=fav, b=btn: self._fav_context_menu(f, b))
            self.fav_layout.insertWidget(self.fav_layout.count() - 1, btn)
        self.fav_frame.setVisible(bool(favs))

    def _use_favorite(self, path: str):
        self.inp_folder.setText(path)
        cfg = load_config()
        cfg["folder"] = path
        save_config(cfg)
        self.status.showMessage(f"Dossier : {path}")

    def _fav_context_menu(self, fav: dict, btn):
        from PyQt6.QtWidgets import QMenu
        menu = QMenu(self)
        menu.setStyleSheet(f"""
            QMenu {{ background:{BG_PANEL}; border:1px solid {BORDER}; color:{TEXT_PRIMARY}; }}
            QMenu::item:selected {{ background:{BG_ROW_SEL}; }}
            QMenu::item {{ padding: 5px 20px; }}
        """)
        rename_action = menu.addAction("✏️  Renommer")
        rename_action.triggered.connect(lambda: self._rename_favorite(fav))
        remove_action = menu.addAction("🗑  Supprimer")
        remove_action.triggered.connect(lambda: self._remove_favorite(fav))
        menu.exec(btn.mapToGlobal(btn.rect().bottomLeft()))

    def _rename_favorite(self, fav: dict):
        from PyQt6.QtWidgets import QInputDialog
        new_name, ok = QInputDialog.getText(
            self, "Renommer le favori", "Nouveau nom :",
            text=fav["name"]
        )
        if ok and new_name.strip():
            favs = load_favorites()
            for f in favs:
                if f["path"] == fav["path"]:
                    f["name"] = new_name.strip()
                    break
            save_favorites(favs)
            self._refresh_favorites()
            self.status.showMessage(f"Favori renommé : {new_name.strip()}")

    def _remove_favorite(self, fav: dict):
        favs = load_favorites()
        favs = [f for f in favs if f["path"] != fav["path"]]
        save_favorites(favs)
        self._refresh_favorites()
        self.status.showMessage("Favori supprimé.")

    # ── Actions ───────────────────────────────────────────────────────────────
    def _browse(self):
        folder = QFileDialog.getExistingDirectory(self, "Sélectionner un dossier")
        if folder:
            self.inp_folder.setText(folder)
            cfg = load_config()
            cfg["folder"] = folder
            save_config(cfg)

    def _clear(self):
        self._model.clear()
        self.status.showMessage("Prêt.")

    def _start_search(self):
        folder = self.inp_folder.text().strip()
        terms_raw = self.inp_terms.text().strip()

        if not folder or not os.path.isdir(folder):
            self.status.showMessage("⚠  Dossier invalide.")
            return
        if not terms_raw:
            self.status.showMessage("⚠  Entrez au moins un terme.")
            return
        if not HAS_DOCX and not HAS_PDF:
            self.status.showMessage("⚠  Installez python-docx et pymupdf.")
            return

        terms, mode = parse_query(terms_raw)
        if not terms:
            return

        cfg = load_config()
        cfg["folder"] = folder
        save_config(cfg)
        self._clear()
        self._search_start = time.perf_counter()
        self.btn_search.setEnabled(False)
        self.btn_stop.setEnabled(True)
        self.progress.show()
        self.progress.setRange(0, 0)

        self._worker = SearchWorker(
            folder, terms,
            self.chk_case.isChecked(),
            self.chk_word.isChecked(),
            self.chk_recurse.isChecked(),
            mode,
        )
        self._worker.result_found.connect(self._on_result)
        self._worker.progress.connect(self._on_progress)
        self._worker.finished.connect(self._on_finished)
        self._worker.start()

    def _stop_search(self):
        if self._worker:
            self._worker.stop()
        self.status.showMessage("Arrêt en cours…")

    def _on_result(self, r: dict):
        self._model.add_result(r)

    def _on_progress(self, idx: int, total: int, filename: str):
        self.status.showMessage(f"Analyse {idx}/{total} : {filename}")

    def _on_finished(self, count: int, total: int):
        self.btn_search.setEnabled(True)
        self.btn_stop.setEnabled(False)
        self.progress.hide()
        elapsed = time.perf_counter() - getattr(self, "_search_start", 0)
        elapsed_str = f"{elapsed:.2f}s" if elapsed >= 1 else f"{elapsed*1000:.0f}ms"
        self.status.showMessage(
            f"{count} occurrence(s) trouvée(s) dans {total} fichier(s) — {elapsed_str}")

    def _open_file(self, index: QModelIndex):
        r = self._model.get_row(index.row())
        path = r.get("file", "")
        page = r.get("page")  # None pour docx, int pour pdf
        ext = os.path.splitext(path)[1].lower()

        if not os.path.exists(path):
            return

        if ext == ".pdf" and page is not None:
            self._open_pdf_at_page(path, page)
        else:
            # DOCX : copier l'extrait brut dans le presse-papier pour Ctrl+F
            to_copy = r.get("ctrlf") or r.get("term", "").split(" + ")[0]
            if to_copy and r.get("term") != "ERREUR":
                QApplication.clipboard().setText(to_copy)
                self.status.showMessage(
                    "Extrait copié — faites Ctrl+F dans Word puis Ctrl+V")
            if os.name == "nt":
                os.startfile(path)
            else:
                os.system(f'xdg-open "{path}"')

    def _open_pdf_at_page(self, path: str, page: int):
        """Ouvre un PDF à une page précise selon le lecteur disponible."""
        if os.name == "nt":
            import subprocess
            # Essayer SumatraPDF en premier (supporte -page)
            sumatra_paths = [
                r"C:\Program Files\SumatraPDF\SumatraPDF.exe",
                r"C:\Program Files (x86)\SumatraPDF\SumatraPDF.exe",
                os.path.expanduser(r"~\AppData\Local\SumatraPDF\SumatraPDF.exe"),
            ]
            for sumatra in sumatra_paths:
                if os.path.exists(sumatra):
                    subprocess.Popen([sumatra, "-page", str(page), path])
                    self.status.showMessage(f"PDF ouvert page {page}")
                    return
            # Essayer Adobe Acrobat (supporte /A page=N)
            acrobat_paths = [
                r"C:\Program Files\Adobe\Acrobat DC\Acrobat\Acrobat.exe",
                r"C:\Program Files (x86)\Adobe\Acrobat DC\Acrobat\Acrobat.exe",
                r"C:\Program Files\Adobe\Acrobat Reader DC\Reader\AcroRd32.exe",
                r"C:\Program Files (x86)\Adobe\Acrobat Reader DC\Reader\AcroRd32.exe",
            ]
            for acrobat in acrobat_paths:
                if os.path.exists(acrobat):
                    subprocess.Popen([acrobat, "/A", f"page={page}", path])
                    self.status.showMessage(f"PDF ouvert page {page}")
                    return
            # Fallback : ouverture normale
            os.startfile(path)
            self.status.showMessage(
                f"Page {page} — installez SumatraPDF pour l'ouverture directe à la page")
        else:
            # Linux : evince ou okular supportent --page-label
            import subprocess
            for reader, arg in [("evince", f"--page-label={page}"),
                                 ("okular", f"--page={page - 1}"),
                                 ("zathura", f"-P {page}")]:
                if subprocess.run(["which", reader], capture_output=True).returncode == 0:
                    subprocess.Popen([reader, arg, path])
                    self.status.showMessage(f"PDF ouvert page {page}")
                    return
            os.system(f'xdg-open "{path}"')

    def _export_csv(self):
        results = self._model.all_results()
        if not results:
            self.status.showMessage("Aucun résultat à exporter.")
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "Enregistrer", "", "CSV (*.csv);;Tous (*.*)")
        if not path:
            return
        import csv
        with open(path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.DictWriter(f, fieldnames=["file", "term", "context"])
            writer.writeheader()
            writer.writerows(results)
        self.status.showMessage(f"Exporté : {path}")


# ── Point d'entrée ────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    app = QApplication(sys.argv)
    app.setStyleSheet(STYLESHEET)
    window = SearchApp()
    window.show()
    sys.exit(app.exec())

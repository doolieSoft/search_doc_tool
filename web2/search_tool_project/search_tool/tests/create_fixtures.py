"""
Script to generate test fixture files programmatically.
Run once: python create_fixtures.py
"""
import os
import zipfile

FIXTURES_DIR = os.path.dirname(__file__) + "/fixtures"


def create_sample_pdf():
    """Create a real multi-page PDF with known text content using fitz."""
    import fitz

    path = os.path.join(FIXTURES_DIR, "sample.pdf")
    doc = fitz.open()

    # Page 1
    page = doc.new_page()
    page.insert_text((72, 100), "Page 1: Bienvenue dans Search Doc Tool", fontsize=14)
    page.insert_text((72, 130), "Ce fichier contient du texte recherchable.", fontsize=11)
    page.insert_text((72, 160), "Terme unique: FIXTURE_TERM_ONE", fontsize=11)

    # Page 2
    page = doc.new_page()
    page.insert_text((72, 100), "Page 2: Deuxième page du document", fontsize=14)
    page.insert_text((72, 130), "Autre contenu: FIXTURE_TERM_TWO", fontsize=11)
    page.insert_text((72, 160), "Accents: éàü café résumé", fontsize=11)

    # Page 3
    page = doc.new_page()
    page.insert_text((72, 100), "Page 3: AND test page", fontsize=14)
    page.insert_text((72, 130), "FIXTURE_TERM_ONE et FIXTURE_TERM_TWO sur la même page", fontsize=11)

    doc.save(path)
    doc.close()
    print(f"Created: {path}")


def create_sample_docx():
    """Create a real DOCX with known text content using python-docx."""
    from docx import Document

    path = os.path.join(FIXTURES_DIR, "sample.docx")
    doc = Document()
    doc.add_heading("Sample DOCX for Testing", level=1)
    doc.add_paragraph("Ce document contient du texte recherchable.")
    doc.add_paragraph("Terme unique: DOCX_TERM_ONE")
    doc.add_paragraph("Autre terme: DOCX_TERM_TWO")
    doc.add_paragraph("Accents: éàü café résumé")
    doc.save(path)
    print(f"Created: {path}")


def create_sample_encrypted_docx():
    """
    Create a fake encrypted DOCX (OLE compound file, not a valid ZIP).
    python-docx cannot create truly password-protected files, so we simulate it
    by creating an invalid ZIP file that mimics what Office produces for encrypted docs.
    """
    path = os.path.join(FIXTURES_DIR, "sample_encrypted.docx")
    # Write a minimal OLE signature — not a valid ZIP, so zipfile.BadZipFile is raised
    ole_signature = b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1" + b"\x00" * 500
    with open(path, "wb") as f:
        f.write(ole_signature)
    print(f"Created: {path}")


if __name__ == "__main__":
    os.makedirs(FIXTURES_DIR, exist_ok=True)
    create_sample_pdf()
    create_sample_docx()
    create_sample_encrypted_docx()
    print("All fixtures created.")
